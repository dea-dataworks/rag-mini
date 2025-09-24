from typing import List, Tuple
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma

from pypdf import PdfReader
from docx import Document as DocxDocument
from collections import defaultdict  
import io, os, re, math, json, time

from utils.helpers import compute_score_stats
from guardrails import evaluate_guardrails, pick_primary_status

# cache BM25 per-vectorstore to avoid re-tokenizing
_BM25_CACHE = {}  # key: id(vs) -> SimpleBM25

_TOKENIZER = re.compile(r"\w+").findall

def _tokenize(text: str):
    return [t.lower() for t in _TOKENIZER(text or "")]

def _doc_key(doc: Document) -> tuple:
    md = doc.metadata or {}
    return (
        md.get("source", ""),
        md.get("page", None),
        (doc.page_content or "")[:64],  # short prefix to keep keys stable
    )

class SimpleBM25:
    def __init__(self, docs: List[Document], k1: float = 1.5, b: float = 0.75):
        self.k1, self.b = k1, b
        self.docs = docs
        self.N = len(docs)
        self.tf = []              # list[dict[token] -> freq]
        self.df = defaultdict(int)
        self.doc_len = []
        for d in docs:
            toks = _tokenize(d.page_content)
            counts = defaultdict(int)
            for t in toks:
                counts[t] += 1
            self.tf.append(counts)
            self.doc_len.append(sum(counts.values()))
            for t in counts.keys():
                self.df[t] += 1
        self.avgdl = (sum(self.doc_len) / self.N) if self.N else 0.0

    def _idf(self, t: str) -> float:
        n = self.df.get(t, 0)
        # slight smoothing to avoid negatives on rare corpora
        return math.log((self.N - n + 0.5) / (n + 0.5) + 1.0)

    def score(self, query: str, top_m: int = 50) -> List[tuple]:
        if not self.docs:
            return []
        q_toks = _tokenize(query)
        scores = [0.0] * self.N
        for i, counts in enumerate(self.tf):
            dl = self.doc_len[i] or 1
            norm = 1 - self.b + self.b * (dl / (self.avgdl or 1))
            s = 0.0
            for t in q_toks:
                f = counts.get(t, 0)
                if not f:
                    continue
                idf = self._idf(t)
                s += idf * ((f * (self.k1 + 1)) / (f + self.k1 * norm))
            scores[i] = s
        idxs = sorted(range(self.N), key=lambda j: scores[j], reverse=True)[:top_m]
        return [(self.docs[j], scores[j]) for j in idxs if scores[j] > 0]

def _rrf_fuse(
    lists: List[List[tuple]],  # each: [(Document, score), ...] in rank order
    rrf_k: int = 60,
    top_k: int = 4,
) -> List[tuple]:
    acc = defaultdict(float)
    pick = {}  # key -> Document (first seen)
    for result_list in lists:
        for rank, (doc, _score) in enumerate(result_list, start=1):
            key = _doc_key(doc)
            acc[key] += 1.0 / (rrf_k + rank)
            if key not in pick:
                pick[key] = doc
    fused = sorted(acc.items(), key=lambda kv: kv[1], reverse=True)[:top_k]
    return [(pick[key], fused_score) for key, fused_score in fused]

def _get_all_docs_from_chroma(vs) -> List[Document]:
    """
    Try to read the full corpus from Chroma.
    Falls back to empty list if the backend doesn’t support a full dump.
    """
    try:
        raw = vs.get(include=["documents", "metadatas", "ids"])
    except Exception:
        try:
            raw = vs._collection.get(include=["documents", "metadatas", "ids"])
        except Exception:
            return []
    docs = []
    for text, meta in zip(raw.get("documents", []) or [], raw.get("metadatas", []) or []):
        docs.append(Document(page_content=text or "", metadata=meta or {}))
    return docs


def _ensure_bm25_for_vs(vs, candidate_docs: List[Document] | None = None) -> SimpleBM25:
    """
    Build or return cached BM25 for this vectorstore.
    If candidate_docs is provided, we build BM25 over that pool (fallback).
    """
    if candidate_docs is not None:
        return SimpleBM25(candidate_docs)
    key = id(vs)
    if key in _BM25_CACHE:
        return _BM25_CACHE[key]
    corpus = _get_all_docs_from_chroma(vs)
    bm25 = SimpleBM25(corpus)
    _BM25_CACHE[key] = bm25
    return bm25

# # ---------- HELPERS (stubs) ----------
def read_txt(file_obj: io.BytesIO) -> str:
    """Decode a .txt upload to plain text (UTF-8 fallback)."""
    try:
        file_obj.seek(0) 
        return file_obj.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""

def read_pdf_pages(file_obj: io.BytesIO, name: str):
    """Extract text per page from a text-based PDF (no OCR)."""
    try:
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        out = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                out.append(Document(
                    page_content=text,
                    metadata={"source": name, "page": i + 1, "ext": "pdf"}
                ))
        return out
    except Exception:
        return []

def read_docx(file_obj: io.BytesIO) -> str:
    """Extract text from a .docx by joining non-empty paragraphs."""
    try:
        file_obj.seek(0)
        buf = file_obj.read()
        doc = DocxDocument(io.BytesIO(buf))
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception:
        return ""   


def files_to_documents(uploaded_files) -> Tuple[List[Document], List[str]]:
    """Turn Streamlit UploadedFile objects into LangChain Documents.
    Returns: (docs, skipped) where skipped is a list of 'filename — reason'.
    """
    docs: List[Document] = []
    skipped: List[str] = []

    for uf in uploaded_files:
        name = uf.name
        ext = os.path.splitext(name)[1].lower()

        if ext == ".txt":
            text = read_txt(uf)
        elif ext == ".docx":
            text = read_docx(uf)
        elif ext == ".pdf":
            pdf_docs = read_pdf_pages(uf, name)
            if pdf_docs:
                docs.extend(pdf_docs)
                continue
            else:
                skipped.append(f"{name} — no extractable text (empty or parse error)")
                continue
        else:
            skipped.append(f"{name} — unsupported file type ({ext})")
            continue  # ignore other types

        if not text or not text.strip():
            skipped.append(f"{name} — no extractable text (empty or parse error)")
            continue  # skip empty files

        # For .txt keep page=1 (legacy), for .docx omit page markers
        meta = {"source": name}
        if ext == ".txt":
            meta.update({"page": 1, "ext": "txt"})
        elif ext == ".docx":
            meta.update({"ext": "docx"})
        else:
            # Fallback: keep at least source
            meta.update({"ext": ext.lstrip(".")})

        docs.append(Document(page_content=text, metadata=meta))


    return docs, skipped

def chunk_documents(docs, size: int, overlap: int):
    """
    Split Documents into overlapping chunks.
    Assign a deterministic chunk_id per (source[, page]) and per-chunk index.
    - size: target characters per chunk (e.g., 800)
    - overlap: shared chars between neighboring chunks (e.g., 120)
    Returns a new list[Document] (each chunk keeps/enriches metadata).
    """
    if not docs:
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", " ", ""],  # coarse → fine
    )

    out = []
    for doc in docs:
        parts = splitter.split_documents([doc])
        # Build stable IDs scoped to this source and (optional) page
        src = (doc.metadata or {}).get("source", "unknown")
        page = (doc.metadata or {}).get("page", None)

        for i, ch in enumerate(parts, start=1):
            md = dict(ch.metadata or {})
            # Keep page exactly as original (None for DOCX; 1 for TXT; actual page for PDF)
            # Deterministic ID: source + optional page + chunk counter
            id_mid = f"p{page}::" if page is not None else ""
            cid = f"{src}::{id_mid}c{i}"

            md["chunk_index"] = i
            md["chunk_id"] = cid
            # Also provide 'id' for broader compatibility with some tooling
            md.setdefault("id", cid)

            ch.metadata = md
            out.append(ch)

    return out

def get_embeddings(model_name: str):
    """Return an OllamaEmbeddings instance."""
    try:
        return OllamaEmbeddings(model=model_name)
    except Exception as e:
        raise RuntimeError(
            f"Could not initialize embeddings for '{model_name}'. "
            f"Make sure you've pulled it with `ollama pull {model_name}`."
        ) from e
    
def build_or_load_vectorstore(chunks, embedding, persist_dir: str):
    """
    Build a Chroma vector store when chunks are provided (embeds & persists).
    If chunks is empty, return a handle to an existing store (no embedding).
    Note: Newer langchain-chroma persists on create; no .persist() method exists.
    """
    if chunks and len(chunks) > 0:
        # Creates/updates a persisted collection immediately.
        return Chroma.from_documents(
            chunks, embedding=embedding, persist_directory=persist_dir
        )
    # Open existing collection (no embed pass)
    return Chroma(embedding_function=embedding, persist_directory=persist_dir)

# --- post-retrieval helpers (UI-agnostic) ---

def normalize_hits(hits):
    """Return list[(Document, score|None)]."""
    out = []
    for item in hits or []:
        if isinstance(item, tuple) and len(item) == 2:
            out.append(item)
        else:
            out.append((item, None))
    return out

def filter_by_score(pairs, threshold: float):
    return [(d, s) for (d, s) in pairs if (isinstance(s, (float, int)) and s >= threshold) or s is None]

def cap_per_source(pairs, cap: int):
    counts, kept = {}, []
    for d, s in pairs:
        src = (d.metadata or {}).get("source", "unknown")
        counts[src] = counts.get(src, 0) + 1
        if counts[src] <= cap:
            kept.append((d, s))
    return kept

def make_chunk_rows(pairs, snippet_len: int = 240):
    """Rows for the Chunk Inspector table (no Streamlit calls)."""
    rows = []
    for i, (doc, score) in enumerate(pairs, start=1):
        meta = doc.metadata or {}
        full = (doc.page_content or "").replace("\n", " ")
        snippet = full[:snippet_len] + ("…" if len(full) > snippet_len else "")
        rows.append({
            "Rank": i,
            "Score": f"{score:.4f}" if isinstance(score, (float, int)) else "—",
            "File": meta.get("source", "unknown"),
            "Page": meta.get("page", ""),
            "Chunk ID": meta.get("chunk_id") or meta.get("id"),
            "Snippet": snippet,
        })
    return rows

def build_citation_tags(docs):
    """Stable, de-duped `source p.X` tags for captions."""
    cited_pairs = []
    for d in docs:
        m = d.metadata or {}
        cited_pairs.append((m.get("source", "unknown"), m.get("page", None)))
    seen = {}
    for src, pg in cited_pairs:
        key = (src, pg)
        if key not in seen:
            seen[key] = None
    ordered = sorted(seen.keys(), key=lambda t: (t[0], t[1] or 0))
    return [f"{src} p.{pg}" if pg else src for src, pg in ordered]

# --- retrieval ---
def retrieve(
    vs,
    query: str,
    k: int,
    mmr_lambda: float = 0.7,
    mode: str = "dense",  # "dense" | "hybrid" | "bm25"  (NEW)
):
    """
    Return top-k relevant chunks with scores.
    Output: list of (Document, score).
    - dense: similarity + MMR (status quo).
    - hybrid: BM25 + Dense fused via RRF.
    - bm25: sparse-only over the corpus (NEW).
    """
    if vs is None:
        return []

    dense_fetch = max(3 * k, 20)
    # Try to get dense results (also used to seed BM25 if corpus dump is unavailable)
    try:
        dense_results = vs.similarity_search_with_score(query, k=dense_fetch)
    except Exception:
        dense_docs = vs.similarity_search(query, k=dense_fetch)
        dense_results = [(d, 0.0) for d in dense_docs]

    # --- Dense path (unchanged) ---
    if mode == "dense":
        retriever = vs.as_retriever(
            search_type="mmr",
            search_kwargs={"k": k, "fetch_k": max(10, 3 * k), "lambda_mult": mmr_lambda},
        )
        mmr_docs = retriever.invoke(query)
        score_map = {d.page_content: s for d, s in dense_results}
        return [(d, score_map.get(d.page_content)) for d in mmr_docs]

    # Get BM25 index over entire corpus or over dense candidates as a fallback
    full_corpus = _get_all_docs_from_chroma(vs)
    bm25_index = _ensure_bm25_for_vs(
        vs,
        candidate_docs=[d for (d, _s) in dense_results] if not full_corpus else None,
    )
    bm25_top = bm25_index.score(query, top_m=dense_fetch)  # [(Document, score)]

    # --- BM25-only path (NEW) ---
    if mode == "bm25":
        return bm25_top[:k]

    # --- Hybrid path (unchanged): fuse BM25 + Dense via RRF ---
    dense_ranked = [(doc, score) for (doc, score) in dense_results]
    fused = _rrf_fuse([bm25_top, dense_ranked], rrf_k=60, top_k=k)
    return fused

# --- load existing store on app start ---
def load_vectorstore_if_exists(embed_model: str, persist_dir: str):
    """Try to open an existing Chroma store; return vs or None."""
    if not os.path.exists(persist_dir):
        return None
    embedding = OllamaEmbeddings(model=embed_model)
    try:
        return Chroma(embedding_function=embedding, persist_directory=persist_dir)
    except Exception:
        return None
    
# ---------- Functions ----------

def build_index_from_files(
    uploaded_files,
    embed_model: str,
    chunk_size: int,
    chunk_overlap: int,
    persist_dir: str,
    embedding_obj=None,  
):
    """
    Orchestrator for index building:
      1) Convert uploaded files -> Documents
      2) Split documents into chunks
      3) Get embeddings
      4) Build/load Chroma vector store and persist it

    Returns:
        vs: the Chroma vector store object
        stats: dict with summary info
            {'num_docs': int, 'num_chunks': int, 'sources': list[str]}
    """

    docs, skipped = files_to_documents(uploaded_files)
    chunks = chunk_documents(docs, size=chunk_size, overlap=chunk_overlap)
    
    embedding = embedding_obj or get_embeddings(embed_model)
    
    vs = build_or_load_vectorstore(
        chunks=chunks,
        embedding=embedding,
        persist_dir=persist_dir,
    )

    stats = {
        "num_docs": len(docs),
        "num_chunks": len(chunks),
        "sources": list({d.metadata.get("source", "unknown") for d in docs}),
        "skipped_files": skipped,
        "per_file": {}
    }

    # --- Extra stats ---
    total_chars = sum(len(ch.page_content or "") for ch in chunks)
    avg_chunk_len = round(total_chars / stats["num_chunks"], 1) if stats["num_chunks"] else 0

    stats["total_chars"] = total_chars
    stats["avg_chunk_len"] = avg_chunk_len

    # Count pages, chunks, and chars per file
    page_count = defaultdict(set)
    chunk_count = defaultdict(int)
    char_count = defaultdict(int)
    for d in chunks:
        src = d.metadata.get("source", "unknown")
        pg = d.metadata.get("page", None)
        if pg:
            page_count[src].add(pg)
        chunk_count[src] += 1
        char_count[src] += len(d.page_content or "")

    for src in stats["sources"]:
        c = chunk_count.get(src, 0)
        total_c = char_count.get(src, 0)
        stats["per_file"][src] = {
            "pages": len(page_count[src]) if src in page_count else 1,
            "chunks": c,
            "chars": total_c,
            "avg_chunk_len": round(total_c / c, 1) if c else 0,
        }


    # Attach persist_dir and write a manifest next to the vectors
    stats["persist_dir"] = persist_dir
    write_manifest(
        persist_dir,
        stats,
        params={
            "embed_model": embed_model,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
        },
    )
    return vs, stats

# --- Index management (fresh dir + manifest) ---
def make_fresh_index_dir(base_dir: str) -> str:
    """
    Create a new subfolder under base_dir for a clean index build.
    Example: rag_store/demo/idx_20250919_121530
    Also clears BM25 cache so hybrid retrieval can’t see stale corpora.
    """
    ts = time.strftime("%Y%m%d_%H%M%S")
    new_dir = os.path.join(base_dir, f"idx_{ts}")
    os.makedirs(new_dir, exist_ok=True)
    _BM25_CACHE.clear()
    return new_dir

def _manifest_path(persist_dir: str) -> str:
    return os.path.join(persist_dir, "manifest.json")

def write_manifest(persist_dir: str, stats: dict, params: dict | None = None):
    payload = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "persist_dir": persist_dir,
        "num_docs": stats.get("num_docs", 0),
        "num_chunks": stats.get("num_chunks", 0),
        "sources": stats.get("sources", []),
        "per_file": stats.get("per_file", {}),
        "total_chars": stats.get("total_chars", 0),
        "avg_chunk_len": stats.get("avg_chunk_len", 0),
        "params": params or {},
    }

    try:
        with open(_manifest_path(persist_dir), "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception:
        pass  # non-fatal

def read_manifest(persist_dir: str) -> dict | None:
    try:
        with open(_manifest_path(persist_dir), "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None
    
# --- Active pointer (per base) & index discovery ---

def _active_pointer_path(base_dir: str) -> str:
    return os.path.join(base_dir, "_ACTIVE.json")

def save_active_pointer(base_dir: str, index_dir: str) -> None:
    """Persist the active index for this base so the app can auto-load on startup."""
    payload = {"active_index_dir": index_dir}
    try:
        with open(_active_pointer_path(base_dir), "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception:
        pass  # non-fatal

def read_active_pointer(base_dir: str) -> str | None:
    """Return the saved active index dir if it exists and is a valid folder."""
    try:
        with open(_active_pointer_path(base_dir), "r", encoding="utf-8") as f:
            data = json.load(f)
        cand = data.get("active_index_dir")
        if cand and os.path.isdir(cand):
            return cand
    except Exception:
        return None
    return None

def list_index_dirs(base_dir: str) -> list[str]:
    """Return absolute paths to subfolders matching idx_YYYYMMDD_HHMMSS under base_dir."""
    if not os.path.isdir(base_dir):
        return []
    out = []
    for name in os.listdir(base_dir):
        p = os.path.join(base_dir, name)
        if os.path.isdir(p) and name.startswith("idx_"):
            out.append(p)
    return sorted(out)

def find_latest_index_dir(base_dir: str) -> str | None:
    """Pick the latest idx_* folder by lexicographic order (timestamped names)."""
    idxs = list_index_dirs(base_dir)
    return idxs[-1] if idxs else None








# --- Role tagging (very light heuristics) ---
_NUM_RE   = re.compile(r"\b\d{1,2}[:/.-]\d{1,2}[:/.-]\d{2,4}\b|\b\d{4}\b|\b\d+(?:\.\d+)?\b", re.IGNORECASE)
_MONTH_RE = re.compile(r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\b", re.IGNORECASE)
_QUOTE_RE = re.compile(r"[\"“”]")

def _infer_chunk_role(text: str, rank: int) -> str:
    """
    Tiny heuristic:
      - first-ranked chunk with definitional cue -> 'definitional'
      - contains numbers/dates/quotes -> 'fact source'
      - otherwise -> 'context'
    """
    t = (text or "").strip()
    head = t.split("\n", 1)[0][:240].lower()

    if rank == 1 and any(k in head for k in [" is ", " are ", " refers to ", " means ", " defined as "]):
        return "definitional"

    if _NUM_RE.search(t) or _MONTH_RE.search(t) or _QUOTE_RE.search(t):
        return "fact source"

    return "context"

# === EXPORT PAYLOAD HELPERS ===
def _make_chunk_records(pairs, snippet_len: int = 300):
    """
    Convert normalized retrieval pairs into export-friendly dicts.
    Each: rank, score, source, page, chunk_id, snippet, short_snippet, role.
    """
    rows = []
    for i, (doc, score) in enumerate(pairs, start=1):
        meta = doc.metadata or {}
        full = (doc.page_content or "").replace("\n", " ")
        snippet = full[:snippet_len] + ("…" if len(full) > snippet_len else "")
        short = full[:160] + ("…" if len(full) > 160 else "")
        role = _infer_chunk_role(full, rank=i)

        rows.append({
            "rank": i,
            "score": float(score) if isinstance(score, (float, int)) else None,
            "source": meta.get("source", "unknown"),
            "page": meta.get("page", None),
            "chunk_id": meta.get("id") or meta.get("chunk_id"),
            "snippet": snippet,
            "short_snippet": short,
            "role": role,
        })
    return rows

def _dedup_citations(docs):
    """Return unique [{'source': ..., 'page': ...}] in stable order."""
    seen = set()
    out = []
    for d in docs or []:
        m = d.metadata or {}
        key = (m.get("source", "unknown"), m.get("page", None))
        if key not in seen:
            seen.add(key)
            out.append({"source": key[0], "page": key[1]})
    # sort for stability (by source, then page if present)
    out.sort(key=lambda r: (r["source"], r["page"] or 0))
    return out

def build_qa_result(
    question: str,
    answer: str,
    docs_used: list,            # sanitized docs actually sent to LLM
    pairs: list,                # [(Document, score)] used for ranking
    meta: dict | None = None,   # {'model': ..., 'top_k': ..., 'retrieval_mode': ...}
    timings: dict | None = None,
    *,
    # --- NEW: guardrail inputs (optional; kept keyword-only for backward compat) ---
    context_text: str = "",
    sanitize_telemetry: dict | None = None,   # e.g., {"chunks_with_drops": 1, "lines_dropped": 3}
    scrubbed_lines: list[str] | None = None,  # from guardrails.scrub_context()
) -> dict:
    """
    Package a single QA turn into a consistent dict for UI/storage/export.
    Adds:
      • 'retrieved_chunks' with role tags (for Why-this-answer)
      • 'guardrail_statuses' + 'guardrail_primary_status' (for banners/exports)
    """
    meta = dict(meta or {})
    meta.setdefault("top_k", len(pairs) if pairs else 0)
    meta.setdefault("timestamp", time.strftime("%Y-%m-%dT%H:%M:%S"))

    txt = (answer or "").strip()
    gist = (txt[:200] + "…") if len(txt) > 200 else txt

    # Build retrieved chunk rows for UI/export
    chunk_rows = _make_chunk_records(pairs or [], snippet_len=300)

    # --- NEW: evaluate guardrails (warn-first; no_context is the only blocker) ---
    statuses = evaluate_guardrails(
        question=question or "",
        context_text=context_text or "",
        docs_used=docs_used or [],
        sanitize_telemetry=sanitize_telemetry or {},
        scrubbed_lines=scrubbed_lines or [],
        answer_text=answer or "",
        min_chars=40,
    )
    primary = pick_primary_status(statuses)

    return {
        "question": question or "",
        "answer": answer or "",
        "answer_gist": gist,
        "citations": _dedup_citations(docs_used),
        "chunks": chunk_rows,
        "retrieved_chunks": chunk_rows,
        "guardrail_statuses": statuses,                 # NEW
        "guardrail_primary_status": primary,            # NEW
        "meta": {
            "model": meta.get("model"),
            "top_k": int(meta.get("top_k", 0)),
            "retrieval_mode": meta.get("retrieval_mode"),
            "timestamp": meta.get("timestamp"),
        },
        "metrics": {
            "scores": compute_score_stats(pairs or []),
            "timings": timings or {},
        },
    }


